"""
VANESSA ALVES - ADVOCACIA TRABALHISTA & FAMÍLIA
Sistema de Gestão Jurídica Profissional
Versão 2.0 - Melhorada e Otimizada
"""

import customtkinter as ctk
import tkinter.messagebox as msg
from tkinter import ttk
import json
import os
import subprocess
import platform
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==================== CONFIGURAÇÕES ====================
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

ARQUIVO = "VanessaAlves_Dados.json"
COR_OURO = "#d4af37"
COR_OURO_ESCURO = "#b8860b"
COR_FUNDO = "#1a1a1a"
COR_CARD = "#2b2b2b"

# ==================== FUNÇÕES DE DADOS ====================
def carregar():
    if os.path.exists(ARQUIVO):
        try:
            with open(ARQUIVO, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            msg.showerror("Erro", f"Erro ao carregar dados: {str(e)}")
            return {"clientes": {}}
    return {"clientes": {}}

def salvar(d):
    try:
        # Backup antes de salvar
        if os.path.exists(ARQUIVO):
            backup = ARQUIVO.replace(".json", "_backup.json")
            with open(ARQUIVO, "r", encoding="utf-8") as f:
                with open(backup, "w", encoding="utf-8") as fb:
                    fb.write(f.read())
        
        with open(ARQUIVO, "w", encoding="utf-8") as f:
            json.dump(d, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        msg.showerror("Erro", f"Erro ao salvar: {str(e)}")
        return False

def abrir_arquivo(caminho):
    """Abre arquivo de forma multiplataforma"""
    try:
        if platform.system() == 'Windows':
            os.startfile(caminho)
        elif platform.system() == 'Darwin':  # macOS
            subprocess.call(['open', caminho])
        else:  # Linux
            subprocess.call(['xdg-open', caminho])
    except Exception as e:
        msg.showerror("Erro", f"Não foi possível abrir o arquivo: {str(e)}")

# ==================== TEMPLATES DE PETIÇÃO ====================
TEMPLATES = {
    "Trabalhista": {
        "titulo": "RECLAMAÇÃO TRABALHISTA",
        "pedidos": [
            "Reconhecimento do vínculo empregatício",
            "Pagamento de verbas rescisórias",
            "FGTS e multa de 40%",
            "Horas extras e adicional noturno",
            "Aviso prévio indenizado",
            "Danos morais"
        ]
    },
    "Família": {
        "titulo": "AÇÃO DE FAMÍLIA",
        "pedidos": [
            "Regulamentação de guarda",
            "Fixação de alimentos",
            "Regime de visitas",
            "Partilha de bens",
            "Reconhecimento de união estável"
        ]
    },
    "Cível": {
        "titulo": "AÇÃO CÍVEL",
        "pedidos": [
            "Indenização por danos materiais",
            "Indenização por danos morais",
            "Cumprimento de obrigação",
            "Ressarcimento de valores"
        ]
    },
    "Previdenciário": {
        "titulo": "AÇÃO PREVIDENCIÁRIA",
        "pedidos": [
            "Concessão de benefício",
            "Revisão de benefício",
            "Restabelecimento de auxílio",
            "Aposentadoria por invalidez"
        ]
    },
    "Consumidor": {
        "titulo": "AÇÃO CONSUMERISTA",
        "pedidos": [
            "Restituição de valores pagos",
            "Danos morais",
            "Cumprimento de oferta",
            "Cancelamento de contrato"
        ]
    }
}

dados = carregar()

# ==================== FUNÇÕES DE EXTRAÇÃO ====================
def extrair_texto_pdf(caminho):
    """Extrai texto de PDF usando PyPDF2"""
    try:
        import PyPDF2
        texto = ""
        with open(caminho, 'rb') as arquivo:
            leitor = PyPDF2.PdfReader(arquivo)
            for pagina in leitor.pages:
                texto += pagina.extract_text() + "\n"
        return texto
    except ImportError:
        return "[PDF_NAO_SUPORTADO]"
    except Exception as e:
        return f"[ERRO ao ler PDF: {str(e)}]"

def extrair_texto_docx(caminho):
    """Extrai texto de DOCX"""
    try:
        doc = Document(caminho)
        texto = "\n".join([paragrafo.text for paragrafo in doc.paragraphs])
        return texto
    except Exception as e:
        return f"[ERRO ao ler DOCX: {str(e)}]"

def extrair_texto_doc(caminho):
    """Extrai texto de DOC usando python-docx ou alternativa"""
    try:
        return extrair_texto_docx(caminho)
    except:
        return "[ERRO: Arquivo .DOC não suportado diretamente. Converta para .DOCX]"

# ==================== JANELA PRINCIPAL ====================
app = ctk.CTk()
app.geometry("1400x900")
app.title("VANESSA ALVES - Sistema de Gestão Jurídica")
app.configure(fg_color=COR_FUNDO)

# ==================== HEADER ELEGANTE ====================
frame_header = ctk.CTkFrame(app, fg_color=COR_CARD, height=120, corner_radius=0)
frame_header.pack(fill="x", pady=0)
frame_header.pack_propagate(False)

# Logo estilizada
frame_logo = ctk.CTkFrame(frame_header, fg_color="transparent")
frame_logo.pack(side="left", padx=40, pady=20)

logo_label = ctk.CTkLabel(
    frame_logo, 
    text="⚖", 
    font=("Segoe UI Emoji", 60), 
    text_color=COR_OURO
)
logo_label.pack()

# Título e subtítulo
frame_titulo = ctk.CTkFrame(frame_header, fg_color="transparent")
frame_titulo.pack(side="left", padx=10, pady=20)

ctk.CTkLabel(
    frame_titulo, 
    text="VANESSA ALVES", 
    font=("Montserrat", 32, "bold"), 
    text_color=COR_OURO
).pack(anchor="w")

ctk.CTkLabel(
    frame_titulo, 
    text="Advocacia Trabalhista & Direito de Família", 
    font=("Montserrat", 14), 
    text_color="#888888"
).pack(anchor="w")

ctk.CTkLabel(
    frame_titulo, 
    text="Sistema de Gestão Jurídica Profissional", 
    font=("Montserrat", 11), 
    text_color="#666666"
).pack(anchor="w", pady=(5, 0))

# Informações do sistema
frame_info = ctk.CTkFrame(frame_header, fg_color="transparent")
frame_info.pack(side="right", padx=40, pady=20)

total_clientes = len(dados["clientes"])
total_processos = sum(len(c["processos"]) for c in dados["clientes"].values())

ctk.CTkLabel(
    frame_info, 
    text=f"📊 {total_clientes} Clientes", 
    font=("Arial", 13), 
    text_color=COR_OURO
).pack(anchor="e")

ctk.CTkLabel(
    frame_info, 
    text=f"⚖️ {total_processos} Processos", 
    font=("Arial", 13), 
    text_color=COR_OURO
).pack(anchor="e")

ctk.CTkLabel(
    frame_info, 
    text=datetime.now().strftime("%d/%m/%Y"), 
    font=("Arial", 11), 
    text_color="#888888"
).pack(anchor="e", pady=(5, 0))

# ==================== TABS PRINCIPAIS ====================
tab = ctk.CTkTabview(app, fg_color=COR_FUNDO, segmented_button_fg_color=COR_CARD,
                     segmented_button_selected_color=COR_OURO_ESCURO,
                     segmented_button_selected_hover_color=COR_OURO)
tab.pack(fill="both", expand=True, padx=20, pady=10)

tab.add("👤 Cliente")
tab.add("📋 Processo")
tab.add("📤 Importar")
tab.add("📄 Petição")
tab.add("📊 Painel")
tab.add("🔍 Buscar")

# ==================== ABA CLIENTE ====================
f1 = tab.tab("👤 Cliente")
f1.configure(fg_color=COR_FUNDO)

# Container com scroll
scroll_cliente = ctk.CTkScrollableFrame(f1, fg_color="transparent")
scroll_cliente.pack(fill="both", expand=True, padx=20, pady=20)

ctk.CTkLabel(
    scroll_cliente, 
    text="CADASTRO DE CLIENTE", 
    font=("Montserrat", 24, "bold"), 
    text_color=COR_OURO
).pack(pady=(0, 30))

# Frame para formulário
frame_form = ctk.CTkFrame(scroll_cliente, fg_color=COR_CARD, corner_radius=15)
frame_form.pack(fill="x", padx=50)

campos_cliente = [
    ("Nome Completo", "text"),
    ("CPF", "text"),
    ("RG", "text"),
    ("Telefone", "text"),
    ("E-mail", "text"),
    ("Endereço", "text"),
    ("Profissão", "text"),
    ("Estado Civil", "combo")
]

entradas_cliente = {}

for i, (campo, tipo) in enumerate(campos_cliente):
    frame_campo = ctk.CTkFrame(frame_form, fg_color="transparent")
    frame_campo.pack(fill="x", padx=40, pady=15)
    
    ctk.CTkLabel(
        frame_campo, 
        text=campo + ":", 
        font=("Arial", 14, "bold"), 
        text_color=COR_OURO
    ).pack(anchor="w")
    
    if tipo == "combo":
        e = ctk.CTkComboBox(
            frame_campo, 
            values=["Solteiro(a)", "Casado(a)", "Divorciado(a)", "Viúvo(a)", "União Estável"],
            width=600, 
            height=45,
            font=("Arial", 13)
        )
        e.set("Solteiro(a)")
    else:
        e = ctk.CTkEntry(
            frame_campo, 
            width=600, 
            height=45,
            font=("Arial", 13),
            placeholder_text=f"Digite o {campo.lower()}"
        )
    e.pack(pady=(5, 0))
    entradas_cliente[campo] = e

# Observações
frame_obs = ctk.CTkFrame(frame_form, fg_color="transparent")
frame_obs.pack(fill="x", padx=40, pady=15)

ctk.CTkLabel(
    frame_obs, 
    text="Observações:", 
    font=("Arial", 14, "bold"), 
    text_color=COR_OURO
).pack(anchor="w")

obs_cliente = ctk.CTkTextbox(frame_obs, width=600, height=100, font=("Arial", 12))
obs_cliente.pack(pady=(5, 0))

def limpar_campos_cliente():
    for campo, entrada in entradas_cliente.items():
        if isinstance(entrada, ctk.CTkEntry):
            entrada.delete(0, "end")
        elif isinstance(entrada, ctk.CTkComboBox):
            entrada.set("Solteiro(a)")
    obs_cliente.delete("1.0", "end")

def salvar_cliente():
    cpf = entradas_cliente["CPF"].get().strip()
    nome = entradas_cliente["Nome Completo"].get().strip()
    
    if not cpf or not nome:
        msg.showerror("Erro", "CPF e Nome são obrigatórios!")
        return
    
    if cpf in dados["clientes"]:
        if not msg.askyesno("Confirmar", "Cliente já existe. Deseja atualizar os dados?"):
            return
    
    dados["clientes"][cpf] = {
        "nome": nome,
        "rg": entradas_cliente["RG"].get().strip(),
        "tel": entradas_cliente["Telefone"].get().strip(),
        "email": entradas_cliente["E-mail"].get().strip(),
        "endereco": entradas_cliente["Endereço"].get().strip(),
        "profissao": entradas_cliente["Profissão"].get().strip(),
        "estado_civil": entradas_cliente["Estado Civil"].get(),
        "observacoes": obs_cliente.get("1.0", "end").strip(),
        "data_cadastro": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "processos": dados["clientes"].get(cpf, {}).get("processos", [])
    }
    
    if salvar(dados):
        msg.showinfo("Sucesso", f"✅ Cliente {nome} cadastrado com sucesso!")
        limpar_campos_cliente()
        atualizar_painel()

# Botões
frame_btns_cliente = ctk.CTkFrame(frame_form, fg_color="transparent")
frame_btns_cliente.pack(pady=30)

ctk.CTkButton(
    frame_btns_cliente,
    text="💾 SALVAR CLIENTE",
    command=salvar_cliente,
    fg_color=COR_OURO_ESCURO,
    hover_color=COR_OURO,
    height=55,
    width=250,
    font=("Arial", 16, "bold"),
    corner_radius=10
).pack(side="left", padx=10)

ctk.CTkButton(
    frame_btns_cliente,
    text="🗑️ LIMPAR",
    command=limpar_campos_cliente,
    fg_color="#8B0000",
    hover_color="#A52A2A",
    height=55,
    width=150,
    font=("Arial", 16, "bold"),
    corner_radius=10
).pack(side="left", padx=10)

# ==================== ABA PROCESSO ====================
f2 = tab.tab("📋 Processo")
f2.configure(fg_color=COR_FUNDO)

scroll_processo = ctk.CTkScrollableFrame(f2, fg_color="transparent")
scroll_processo.pack(fill="both", expand=True, padx=20, pady=20)

ctk.CTkLabel(
    scroll_processo,
    text="CADASTRO DE PROCESSO",
    font=("Montserrat", 24, "bold"),
    text_color=COR_OURO
).pack(pady=(0, 30))

frame_proc = ctk.CTkFrame(scroll_processo, fg_color=COR_CARD, corner_radius=15)
frame_proc.pack(fill="x", padx=50)

# CPF do Cliente
frame_cpf_proc = ctk.CTkFrame(frame_proc, fg_color="transparent")
frame_cpf_proc.pack(fill="x", padx=40, pady=20)

ctk.CTkLabel(
    frame_cpf_proc,
    text="CPF do Cliente:",
    font=("Arial", 14, "bold"),
    text_color=COR_OURO
).pack(anchor="w")

cpf_proc = ctk.CTkEntry(frame_cpf_proc, width=400, height=45, font=("Arial", 13))
cpf_proc.pack(pady=(5, 0), anchor="w")

nome_cliente_label = ctk.CTkLabel(
    frame_cpf_proc,
    text="",
    font=("Arial", 12),
    text_color="#00ff00"
)
nome_cliente_label.pack(anchor="w", pady=(5, 0))

def verificar_cliente(*args):
    cpf = cpf_proc.get().strip()
    if cpf in dados["clientes"]:
        nome_cliente_label.configure(text=f"✓ Cliente: {dados['clientes'][cpf]['nome']}")
    else:
        nome_cliente_label.configure(text="✗ Cliente não encontrado")

cpf_proc.bind("<KeyRelease>", verificar_cliente)

# Dados do Processo
campos_processo = [
    ("Número do Processo", "text"),
    ("Comarca", "text"),
    ("Vara", "text"),
    ("Valor da Causa (R$)", "text")
]

entradas_processo = {}

for campo, tipo in campos_processo:
    frame_campo = ctk.CTkFrame(frame_proc, fg_color="transparent")
    frame_campo.pack(fill="x", padx=40, pady=15)
    
    ctk.CTkLabel(
        frame_campo,
        text=campo + ":",
        font=("Arial", 14, "bold"),
        text_color=COR_OURO
    ).pack(anchor="w")
    
    e = ctk.CTkEntry(
        frame_campo,
        width=600,
        height=45,
        font=("Arial", 13)
    )
    e.pack(pady=(5, 0))
    entradas_processo[campo] = e

# Área do Direito
frame_area = ctk.CTkFrame(frame_proc, fg_color="transparent")
frame_area.pack(fill="x", padx=40, pady=15)

ctk.CTkLabel(
    frame_area,
    text="Área do Direito:",
    font=("Arial", 14, "bold"),
    text_color=COR_OURO
).pack(anchor="w")

area_proc = ctk.CTkComboBox(
    frame_area,
    values=list(TEMPLATES.keys()),
    width=600,
    height=45,
    font=("Arial", 13)
)
area_proc.set("Trabalhista")
area_proc.pack(pady=(5, 0))

# Resumo do caso
frame_resumo = ctk.CTkFrame(frame_proc, fg_color="transparent")
frame_resumo.pack(fill="x", padx=40, pady=15)

ctk.CTkLabel(
    frame_resumo,
    text="Resumo do Caso:",
    font=("Arial", 14, "bold"),
    text_color=COR_OURO
).pack(anchor="w")

resumo_proc = ctk.CTkTextbox(frame_resumo, width=600, height=200, font=("Arial", 12))
resumo_proc.pack(pady=(5, 0))

def limpar_campos_processo():
    cpf_proc.delete(0, "end")
    nome_cliente_label.configure(text="")
    for entrada in entradas_processo.values():
        entrada.delete(0, "end")
    area_proc.set("Trabalhista")
    resumo_proc.delete("1.0", "end")

def salvar_processo():
    cpf = cpf_proc.get().strip()
    
    if cpf not in dados["clientes"]:
        msg.showerror("Erro", "Cliente não encontrado! Cadastre o cliente primeiro.")
        return
    
    num = entradas_processo["Número do Processo"].get().strip()
    if not num:
        msg.showerror("Erro", "Número do processo é obrigatório!")
        return
    
    # Verifica se processo já existe
    for proc in dados["clientes"][cpf]["processos"]:
        if proc["numero"] == num:
            msg.showerror("Erro", "Este número de processo já está cadastrado para este cliente!")
            return
    
    processo = {
        "numero": num,
        "comarca": entradas_processo["Comarca"].get().strip(),
        "vara": entradas_processo["Vara"].get().strip(),
        "valor_causa": entradas_processo["Valor da Causa (R$)"].get().strip(),
        "area": area_proc.get(),
        "resumo": resumo_proc.get("1.0", "end").strip(),
        "data_cadastro": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "status": "Em andamento"
    }
    
    dados["clientes"][cpf]["processos"].append(processo)
    
    if salvar(dados):
        msg.showinfo("Sucesso", f"✅ Processo {num} cadastrado com sucesso!")
        limpar_campos_processo()
        atualizar_painel()

# Botões
frame_btns_proc = ctk.CTkFrame(frame_proc, fg_color="transparent")
frame_btns_proc.pack(pady=30)

ctk.CTkButton(
    frame_btns_proc,
    text="💾 SALVAR PROCESSO",
    command=salvar_processo,
    fg_color="#228B22",
    hover_color="#32CD32",
    height=55,
    width=250,
    font=("Arial", 16, "bold"),
    corner_radius=10
).pack(side="left", padx=10)

ctk.CTkButton(
    frame_btns_proc,
    text="🗑️ LIMPAR",
    command=limpar_campos_processo,
    fg_color="#8B0000",
    hover_color="#A52A2A",
    height=55,
    width=150,
    font=("Arial", 16, "bold"),
    corner_radius=10
).pack(side="left", padx=10)

# ==================== ABA IMPORTAR ====================
f_import = tab.tab("📤 Importar")
f_import.configure(fg_color=COR_FUNDO)

scroll_import = ctk.CTkScrollableFrame(f_import, fg_color="transparent")
scroll_import.pack(fill="both", expand=True, padx=20, pady=20)

ctk.CTkLabel(
    scroll_import,
    text="IMPORTAÇÃO AUTOMÁTICA DE DOCUMENTOS",
    font=("Montserrat", 24, "bold"),
    text_color=COR_OURO
).pack(pady=(0, 10))

ctk.CTkLabel(
    scroll_import,
    text="Arraste arquivos PDF, DOCX ou DOC para extração automática de dados",
    font=("Arial", 12),
    text_color="#888888"
).pack(pady=(0, 30))

frame_import = ctk.CTkFrame(scroll_import, fg_color=COR_CARD, corner_radius=15)
frame_import.pack(fill="both", expand=True, padx=50)

# Frame para seleção de arquivos
frame_select_files = ctk.CTkFrame(frame_import, fg_color="transparent")
frame_select_files.pack(fill="x", padx=40, pady=30)

ctk.CTkLabel(
    frame_select_files,
    text="📁 Selecione os arquivos para importar:",
    font=("Arial", 14, "bold"),
    text_color=COR_OURO
).pack(anchor="w", pady=(0, 10))

# Lista de arquivos selecionados
frame_files_list = ctk.CTkFrame(frame_select_files, fg_color="#1a1a1a", corner_radius=10)
frame_files_list.pack(fill="both", expand=True, pady=10)

text_files = ctk.CTkTextbox(frame_files_list, height=150, font=("Arial", 11))
text_files.pack(fill="both", expand=True, padx=10, pady=10)
text_files.insert("1.0", "Nenhum arquivo selecionado ainda...\n\n✅ Formatos aceitos:\n📝 DOCX, DOC (100% funcional)\n📄 PDF (requer: python -m pip install PyPDF2)")

arquivos_selecionados = []

def selecionar_arquivos():
    from tkinter import filedialog
    arquivos = filedialog.askopenfilenames(
        title="Selecione os documentos",
        filetypes=[
            ("Word (recomendado)", "*.docx *.doc"),
            ("PDF (requer PyPDF2)", "*.pdf"),
            ("Todos suportados", "*.pdf *.docx *.doc"),
            ("Todos", "*.*")
        ]
    )
    
    if arquivos:
        arquivos_selecionados.clear()
        arquivos_selecionados.extend(arquivos)
        
        text_files.delete("1.0", "end")
        text_files.insert("1.0", f"✅ {len(arquivos)} arquivo(s) selecionado(s):\n\n")
        for arq in arquivos:
            nome_arq = os.path.basename(arq)
            extensao = arq.lower().split('.')[-1]
            icone = "📄" if extensao == "pdf" else "📝"
            text_files.insert("end", f"{icone} {nome_arq}\n")

frame_btns_select = ctk.CTkFrame(frame_select_files, fg_color="transparent")
frame_btns_select.pack(pady=10)

ctk.CTkButton(
    frame_btns_select,
    text="📁 SELECIONAR ARQUIVOS",
    command=selecionar_arquivos,
    fg_color=COR_OURO_ESCURO,
    hover_color=COR_OURO,
    height=50,
    width=250,
    font=("Arial", 14, "bold"),
    corner_radius=10
).pack(side="left", padx=5)

# Aviso sobre detecção automática
frame_aviso = ctk.CTkFrame(frame_import, fg_color="#1a4d1a", corner_radius=10)
frame_aviso.pack(fill="x", padx=40, pady=20)

ctk.CTkLabel(
    frame_aviso,
    text="ℹ️ O sistema detectará automaticamente os nomes das partes no documento",
    font=("Arial", 13),
    text_color="#90EE90"
).pack(padx=15, pady=12)

# Área de resultados da extração
frame_resultados = ctk.CTkFrame(frame_import, fg_color="transparent")
frame_resultados.pack(fill="both", expand=True, padx=40, pady=20)

ctk.CTkLabel(
    frame_resultados,
    text="📊 Dados Extraídos:",
    font=("Arial", 14, "bold"),
    text_color=COR_OURO
).pack(anchor="w", pady=(0, 10))

text_resultados = ctk.CTkTextbox(frame_resultados, height=300, font=("Arial", 11))
text_resultados.pack(fill="both", expand=True)
text_resultados.insert("1.0", "Os dados extraídos aparecerão aqui após o processamento...")

def processar_importacao():
    if not arquivos_selecionados:
        msg.showerror("Erro", "Selecione pelo menos um arquivo!")
        return
    
    text_resultados.delete("1.0", "end")
    text_resultados.insert("end", "🔄 Processando arquivos e detectando partes...\n\n")
    app.update()
    
    processos_importados = 0
    erros = 0
    
    for arquivo in arquivos_selecionados:
        nome_arquivo = os.path.basename(arquivo)
        extensao = arquivo.lower().split('.')[-1]
        
        text_resultados.insert("end", f"\n{'='*60}\n")
        text_resultados.insert("end", f"📄 Processando: {nome_arquivo}\n")
        text_resultados.insert("end", f"{'='*60}\n\n")
        app.update()
        
        # Extrair texto baseado na extensão
        if extensao == 'pdf':
            texto = extrair_texto_pdf(arquivo)
            if texto == "[PDF_NAO_SUPORTADO]":
                text_resultados.insert("end", "⚠️ PDFs não suportados (PyPDF2 não instalado)\n")
                text_resultados.insert("end", "💡 SOLUÇÃO: Use arquivos DOCX ou instale PyPDF2 com:\n")
                text_resultados.insert("end", "   python -m pip install PyPDF2\n\n")
                erros += 1
                continue
        elif extensao == 'docx':
            texto = extrair_texto_docx(arquivo)
        elif extensao == 'doc':
            texto = extrair_texto_doc(arquivo)
        else:
            texto = None
            text_resultados.insert("end", f"❌ Formato não suportado: {extensao}\n")
            erros += 1
            continue
        
        if not texto or "[ERRO" in texto:
            text_resultados.insert("end", f"❌ Erro ao extrair texto: {texto}\n")
            erros += 1
            continue
        
        # ========== EXTRAÇÃO DE NOMES DAS PARTES ==========
        import re
        
        nomes_encontrados = set()
        
        # Padrão 1: Nomes após palavras-chave de identificação
        padroes = [
            r'(?:Autor|Autora|Requerente|Reclamante|Cliente|Impetrante)[\s:]+([A-ZÀÁÂÃÄÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ][a-zàáâãäçèéêëìíîïñòóôõöùúûü]+(?:\s+(?:da|de|do|dos|das|e)\s+)?(?:[A-ZÀÁÂÃÄÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ][a-zàáâãäçèéêëìíîïñòóôõöùúûü]+)+)',
            r'(?:Réu|Ré|Requerido|Requerida|Reclamado|Reclamada)[\s:]+([A-ZÀÁÂÃÄÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ][a-zàáâãäçèéêëìíîïñòóôõöùúûü]+(?:\s+(?:da|de|do|dos|das|e)\s+)?(?:[A-ZÀÁÂÃÄÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ][a-zàáâãäçèéêëìíîïñòóôõöùúûü]+)+)',
            r'([A-ZÀÁÂÃÄÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ][a-zàáâãäçèéêëìíîïñòóôõöùúûü]+(?:\s+(?:da|de|do|dos|das|e)\s+)?(?:[A-ZÀÁÂÃÄÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ][a-zàáâãäçèéêëìíîïñòóôõöùúûü]+)+),?\s+(?:brasileiro|brasileira|portador|portadora|inscrito|inscrita|CPF|RG)',
            r'(?:Nome|Cliente)[\s:]+([A-ZÀÁÂÃÄÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ][a-zàáâãäçèéêëìíîïñòóôõöùúûü]+(?:\s+(?:da|de|do|dos|das|e)\s+)?(?:[A-ZÀÁÂÃÄÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ][a-zàáâãäçèéêëìíîïñòóôõöùúûü]+)+)',
        ]
        
        for padrao in padroes:
            matches = re.finditer(padrao, texto, re.MULTILINE | re.IGNORECASE)
            for match in matches:
                nome = match.group(1).strip()
                if len(nome) > 8 and nome not in ['Estado', 'União', 'Município', 'Comarca', 'Brasil']:
                    nomes_encontrados.add(nome)
        
        # Padrão 2: Empresas (Ltd, LTDA, S.A., etc)
        empresas = re.findall(
            r'([A-ZÀÁÂÃÄÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜ][A-Za-zàáâãäçèéêëìíîïñòóôõöùúûü\s&]+(?:LTDA|Ltd|S\.?A\.?|ME|MEI|EIRELI|EPP)\.?)',
            texto
        )
        for empresa in empresas:
            if len(empresa) > 5:
                nomes_encontrados.add(empresa.strip())
        
        lista_nomes = list(nomes_encontrados)
        
        text_resultados.insert("end", f"👥 Nomes detectados: {len(lista_nomes)}\n\n")
        
        if not lista_nomes:
            text_resultados.insert("end", "⚠️ Nenhum nome detectado automaticamente.\n")
            text_resultados.insert("end", "Por favor, digite o nome do cliente manualmente.\n\n")
            
            janela_nome = ctk.CTkToplevel(app)
            janela_nome.title("Nome do Cliente")
            janela_nome.geometry("500x200")
            janela_nome.grab_set()
            
            ctk.CTkLabel(
                janela_nome,
                text="Nenhum nome foi detectado no documento.\nDigite o nome do cliente:",
                font=("Arial", 14)
            ).pack(pady=20)
            
            entry_nome_manual = ctk.CTkEntry(janela_nome, width=400, height=40)
            entry_nome_manual.pack(pady=10)
            entry_nome_manual.focus()
            
            nome_selecionado = [None]
            
            def confirmar_manual():
                nome_selecionado[0] = entry_nome_manual.get().strip()
                janela_nome.destroy()
            
            ctk.CTkButton(
                janela_nome,
                text="Confirmar",
                command=confirmar_manual,
                fg_color=COR_OURO_ESCURO,
                height=40
            ).pack(pady=10)
            
            janela_nome.wait_window()
            
            if not nome_selecionado[0]:
                text_resultados.insert("end", "❌ Importação cancelada (sem nome).\n")
                erros += 1
                continue
            
            cliente_nome = nome_selecionado[0]
            
        else:
            janela_selecao = ctk.CTkToplevel(app)
            janela_selecao.title("Selecione o Cliente")
            janela_selecao.geometry("700x500")
            janela_selecao.grab_set()
            
            ctk.CTkLabel(
                janela_selecao,
                text=f"📄 Arquivo: {nome_arquivo}",
                font=("Arial", 12, "bold")
            ).pack(pady=10)
            
            ctk.CTkLabel(
                janela_selecao,
                text="Selecione quem é o CLIENTE neste processo:",
                font=("Arial", 16, "bold"),
                text_color=COR_OURO
            ).pack(pady=10)
            
            frame_nomes = ctk.CTkScrollableFrame(janela_selecao, width=600, height=300)
            frame_nomes.pack(pady=10, padx=20, fill="both", expand=True)
            
            var_selecionado = ctk.StringVar(value="")
            
            for nome in sorted(lista_nomes):
                frame_radio = ctk.CTkFrame(frame_nomes, fg_color="#2b2b2b")
                frame_radio.pack(fill="x", pady=5, padx=10)
                
                ctk.CTkRadioButton(
                    frame_radio,
                    text=f"👤 {nome}",
                    variable=var_selecionado,
                    value=nome,
                    font=("Arial", 14),
                    radiobutton_width=20,
                    radiobutton_height=20
                ).pack(anchor="w", padx=15, pady=10)
            
            frame_manual = ctk.CTkFrame(frame_nomes, fg_color="#1a4d1a")
            frame_manual.pack(fill="x", pady=5, padx=10)
            
            ctk.CTkLabel(
                frame_manual,
                text="✍️ Ou digite outro nome:",
                font=("Arial", 12, "bold")
            ).pack(anchor="w", padx=15, pady=(10, 5))
            
            entry_outro = ctk.CTkEntry(frame_manual, width=500, height=35)
            entry_outro.pack(padx=15, pady=(0, 10))
            
            nome_selecionado = [None]
            
            def confirmar_selecao():
                if entry_outro.get().strip():
                    nome_selecionado[0] = entry_outro.get().strip()
                elif var_selecionado.get():
                    nome_selecionado[0] = var_selecionado.get()
                else:
                    msg.showerror("Erro", "Selecione ou digite um nome!")
                    return
                janela_selecao.destroy()
            
            def cancelar_selecao():
                nome_selecionado[0] = None
                janela_selecao.destroy()
            
            frame_btns = ctk.CTkFrame(janela_selecao, fg_color="transparent")
            frame_btns.pack(pady=15)
            
            ctk.CTkButton(
                frame_btns,
                text="✅ Confirmar",
                command=confirmar_selecao,
                fg_color=COR_OURO_ESCURO,
                hover_color=COR_OURO,
                width=150,
                height=40,
                font=("Arial", 14, "bold")
            ).pack(side="left", padx=5)
            
            ctk.CTkButton(
                frame_btns,
                text="❌ Cancelar",
                command=cancelar_selecao,
                fg_color="#8B0000",
                hover_color="#A52A2A",
                width=150,
                height=40,
                font=("Arial", 14, "bold")
            ).pack(side="left", padx=5)
            
            janela_selecao.wait_window()
            
            if not nome_selecionado[0]:
                text_resultados.insert("end", "❌ Importação cancelada pelo usuário.\n")
                erros += 1
                continue
            
            cliente_nome = nome_selecionado[0]
        
        text_resultados.insert("end", f"✅ Cliente selecionado: {cliente_nome}\n\n")
        app.update()
        
        # Buscar CPF ou criar cliente
        cpf_encontrado = None
        for cpf, cliente in dados["clientes"].items():
            if cliente["nome"].lower() == cliente_nome.lower():
                cpf_encontrado = cpf
                break
        
        if not cpf_encontrado:
            janela_cpf = ctk.CTkToplevel(app)
            janela_cpf.title("Cadastrar Novo Cliente")
            janela_cpf.geometry("500x350")
            janela_cpf.grab_set()
            
            ctk.CTkLabel(
                janela_cpf,
                text=f"Cliente '{cliente_nome}' não encontrado no sistema.",
                font=("Arial", 13),
                text_color="#ffaa00"
            ).pack(pady=10)
            
            ctk.CTkLabel(
                janela_cpf,
                text="Digite o CPF do cliente:",
                font=("Arial", 14, "bold"),
                text_color=COR_OURO
            ).pack(pady=5)
            
            entry_cpf = ctk.CTkEntry(janela_cpf, width=300, height=40, font=("Arial", 13))
            entry_cpf.pack(pady=10)
            entry_cpf.focus()
            
            ctk.CTkLabel(
                janela_cpf,
                text="Telefone (opcional):",
                font=("Arial", 12)
            ).pack(pady=5)
            
            entry_tel = ctk.CTkEntry(janela_cpf, width=300, height=35)
            entry_tel.pack(pady=5)
            
            ctk.CTkLabel(
                janela_cpf,
                text="Email (opcional):",
                font=("Arial", 12)
            ).pack(pady=5)
            
            entry_email = ctk.CTkEntry(janela_cpf, width=300, height=35)
            entry_email.pack(pady=5)
            
            cpf_novo = [None]
            
            def salvar_novo_cliente():
                cpf = entry_cpf.get().strip()
                if not cpf:
                    msg.showerror("Erro", "CPF é obrigatório!")
                    return
                
                if cpf in dados["clientes"]:
                    msg.showerror("Erro", "Este CPF já está cadastrado!")
                    return
                
                dados["clientes"][cpf] = {
                    "nome": cliente_nome,
                    "rg": "",
                    "tel": entry_tel.get().strip(),
                    "email": entry_email.get().strip(),
                    "endereco": "",
                    "profissao": "",
                    "estado_civil": "Solteiro(a)",
                    "observacoes": f"Cliente criado via importação automática de {nome_arquivo}",
                    "data_cadastro": datetime.now().strftime("%d/%m/%Y %H:%M"),
                    "processos": []
                }
                
                # SALVAR IMEDIATAMENTE
                if not salvar(dados):
                    msg.showerror("Erro", "Não foi possível salvar o cliente!")
                    return
                
                cpf_novo[0] = cpf
                msg.showinfo("Sucesso", f"✅ Cliente {cliente_nome} cadastrado com sucesso!")
                janela_cpf.destroy()
            
            def cancelar_cadastro():
                cpf_novo[0] = None
                janela_cpf.destroy()
            
            frame_btns_cpf = ctk.CTkFrame(janela_cpf, fg_color="transparent")
            frame_btns_cpf.pack(pady=20)
            
            ctk.CTkButton(
                frame_btns_cpf,
                text="💾 Criar Cliente",
                command=salvar_novo_cliente,
                fg_color=COR_OURO_ESCURO,
                hover_color=COR_OURO,
                height=45,
                width=150,
                font=("Arial", 14, "bold")
            ).pack(side="left", padx=5)
            
            ctk.CTkButton(
                frame_btns_cpf,
                text="❌ Cancelar",
                command=cancelar_cadastro,
                fg_color="#8B0000",
                hover_color="#A52A2A",
                height=45,
                width=150,
                font=("Arial", 14, "bold")
            ).pack(side="left", padx=5)
            
            janela_cpf.wait_window()
            
            if not cpf_novo[0]:
                text_resultados.insert("end", "❌ Cadastro cancelado pelo usuário.\n")
                erros += 1
                continue
            
            cpf_encontrado = cpf_novo[0]
            text_resultados.insert("end", f"✅ Novo cliente cadastrado!\n")
            text_resultados.insert("end", f"   👤 Nome: {cliente_nome}\n")
            text_resultados.insert("end", f"   🆔 CPF: {cpf_encontrado}\n\n")
            app.update()
        
        # Extrair informações do processo
        numero_processo = re.search(r'\d{7}-?\d{2}\.?\d{4}\.?\d{1}\.?\d{2}\.?\d{4}', texto)
        
        if not numero_processo:
            numero_processo = re.search(r'(?:Processo|Proc\.|nº|n°|número)[\s:]+(\d[\d\.\-/]+)', texto, re.IGNORECASE)
        
        area = "Cível"
        texto_lower = texto.lower()
        
        if any(palavra in texto_lower for palavra in ['trabalhista', 'reclamação trabalhista', 'vínculo empregatício', 'fgts', 'ctps']):
            area = "Trabalhista"
        elif any(palavra in texto_lower for palavra in ['família', 'divórcio', 'alimentos', 'guarda', 'pensão alimentícia']):
            area = "Família"
        elif any(palavra in texto_lower for palavra in ['previdenciário', 'aposentadoria', 'benefício', 'inss']):
            area = "Previdenciário"
        elif any(palavra in texto_lower for palavra in ['consumidor', 'consumerista', 'cdc', 'relação de consumo']):
            area = "Consumidor"
        
        valor_causa = re.search(r'[Vv]alor\s+da\s+[Cc]ausa[\s:]+R\$\s*([\d.,]+)', texto)
        comarca = re.search(r'[Cc]omarca[\s:]+([A-ZÀ-Ú][a-zà-ú\s]+)', texto)
        
        processo = {
            "numero": numero_processo.group(1) if numero_processo and len(numero_processo.groups()) > 0 else numero_processo.group(0) if numero_processo else f"IMP-{datetime.now().strftime('%Y%m%d%H%M%S')}",
            "comarca": comarca.group(1).strip() if comarca else "",
            "vara": "",
            "valor_causa": valor_causa.group(1) if valor_causa else "",
            "area": area,
            "resumo": f"📄 Processo importado de: {nome_arquivo}\n\n{texto[:800]}...",
            "data_cadastro": datetime.now().strftime("%d/%m/%Y %H:%M"),
            "status": "Importado - Revisar dados",
            "arquivo_original": arquivo
        }
        
        ja_existe = False
        for proc in dados["clientes"][cpf_encontrado]["processos"]:
            if proc["numero"] == processo["numero"]:
                ja_existe = True
                break
        
        if ja_existe:
            text_resultados.insert("end", f"⚠️ Processo {processo['numero']} já existe para este cliente!\n")
        else:
            dados["clientes"][cpf_encontrado]["processos"].append(processo)
            processos_importados += 1
            
            text_resultados.insert("end", f"✅ PROCESSO IMPORTADO COM SUCESSO!\n")
            text_resultados.insert("end", f"   📋 Número: {processo['numero']}\n")
            text_resultados.insert("end", f"   👤 Cliente: {cliente_nome}\n")
            text_resultados.insert("end", f"   ⚖️ Área: {processo['area']}\n")
            if processo['comarca']:
                text_resultados.insert("end", f"   📍 Comarca: {processo['comarca']}\n")
            if processo['valor_causa']:
                text_resultados.insert("end", f"   💰 Valor: R$ {processo['valor_causa']}\n")
        
        app.update()
    
    if processos_importados > 0:
        if salvar(dados):
            text_resultados.insert("end", f"\n\n{'='*60}\n")
            text_resultados.insert("end", f"🎉 IMPORTAÇÃO CONCLUÍDA COM SUCESSO!\n")
            text_resultados.insert("end", f"   📊 {processos_importados} processo(s) importado(s)\n")
            if erros > 0:
                text_resultados.insert("end", f"   ⚠️ {erros} arquivo(s) com erro\n")
            text_resultados.insert("end", f"{'='*60}\n")
            
            msg.showinfo(
                "Sucesso!",
                f"🎉 {processos_importados} processo(s) importado(s) com sucesso!\n\n"
                f"⚠️ IMPORTANTE: Revise os dados importados na aba 'Painel'"
            )
            atualizar_painel()
            
            arquivos_selecionados.clear()
            text_files.delete("1.0", "end")
            text_files.insert("1.0", "Nenhum arquivo selecionado ainda...\n\n✅ Formatos aceitos:\n📝 DOCX, DOC (100% funcional)\n📄 PDF (requer: python -m pip install PyPDF2)")
    else:
        msg.showwarning("Atenção", "Nenhum processo foi importado. Verifique os arquivos ou os logs.")

ctk.CTkButton(
    frame_import,
    text="🚀 PROCESSAR E IMPORTAR",
    command=processar_importacao,
    fg_color="#228B22",
    hover_color="#32CD32",
    height=60,
    width=350,
    font=("Arial", 16, "bold"),
    corner_radius=15
).pack(pady=30)

# Instruções
frame_instrucoes = ctk.CTkFrame(scroll_import, fg_color=COR_CARD, corner_radius=15)
frame_instrucoes.pack(fill="x", padx=50, pady=20)

ctk.CTkLabel(
    frame_instrucoes,
    text="ℹ️ INSTRUÇÕES DE USO",
    font=("Arial", 14, "bold"),
    text_color=COR_OURO
).pack(anchor="w", padx=20, pady=(15, 10))

instrucoes_texto = """1. Clique em "SELECIONAR ARQUIVOS" e escolha os DOCXs dos processos
2. Clique em "PROCESSAR E IMPORTAR"
3. O sistema detectará automaticamente os nomes das partes envolvidas
4. Selecione qual pessoa/empresa é o SEU CLIENTE
5. Se o cliente não estiver cadastrado, será criado automaticamente
6. O sistema extrairá:
   • Número do processo
   • Área do direito (Trabalhista, Família, etc)
   • Comarca
   • Valor da causa
   • Conteúdo completo

⚠️ IMPORTANTE:
• VOCÊ NÃO PRECISA DIGITAR CPF - o sistema detecta os nomes automaticamente!
• Use arquivos DOCX para melhor compatibilidade
• Para PDFs: instale PyPDF2 com: python -m pip install PyPDF2
• Para .DOC antigos, converta para .DOCX primeiro
• O sistema busca nomes após palavras como: Autor, Requerente, Réu, Cliente, etc
"""

ctk.CTkLabel(
    frame_instrucoes,
    text=instrucoes_texto,
    font=("Arial", 11),
    text_color="#CCCCCC",
    justify="left"
).pack(anchor="w", padx=20, pady=(0, 15))

# ==================== ABA PETIÇÃO ====================
f3 = tab.tab("📄 Petição")
f3.configure(fg_color=COR_FUNDO)

scroll_peticao = ctk.CTkScrollableFrame(f3, fg_color="transparent")
scroll_peticao.pack(fill="both", expand=True, padx=20, pady=20)

ctk.CTkLabel(
    scroll_peticao,
    text="GERAÇÃO DE PETIÇÃO INICIAL",
    font=("Montserrat", 24, "bold"),
    text_color=COR_OURO
).pack(pady=(0, 30))

frame_pet = ctk.CTkFrame(scroll_peticao, fg_color=COR_CARD, corner_radius=15)
frame_pet.pack(fill="x", padx=50)

frame_sel = ctk.CTkFrame(frame_pet, fg_color="transparent")
frame_sel.pack(fill="x", padx=40, pady=20)

ctk.CTkLabel(
    frame_sel,
    text="CPF do Cliente:",
    font=("Arial", 14, "bold"),
    text_color=COR_OURO
).pack(anchor="w")

cpf_pet = ctk.CTkEntry(frame_sel, width=400, height=45, font=("Arial", 13))
cpf_pet.pack(pady=(5, 10), anchor="w")

ctk.CTkLabel(
    frame_sel,
    text="Número do Processo:",
    font=("Arial", 14, "bold"),
    text_color=COR_OURO
).pack(anchor="w")

num_pet = ctk.CTkEntry(frame_sel, width=400, height=45, font=("Arial", 13))
num_pet.pack(pady=(5, 0), anchor="w")

def gerar_peticao():
    cpf = cpf_pet.get().strip()
    
    if cpf not in dados["clientes"]:
        msg.showerror("Erro", "Cliente não encontrado!")
        return
    
    cliente = dados["clientes"][cpf]
    proc = None
    
    for p in cliente["processos"]:
        if p["numero"] == num_pet.get().strip():
            proc = p
            break
    
    if not proc:
        msg.showerror("Erro", "Processo não encontrado!")
        return
    
    try:
        doc = Document()
        
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        
        cabecalho = doc.add_heading("EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(ÍZA) DE DIREITO", 0)
        cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Processo nº: {proc['numero']}")
        doc.add_paragraph(f"Comarca: {proc.get('comarca', 'A definir')}")
        doc.add_paragraph(f"Vara: {proc.get('vara', 'A definir')}")
        doc.add_paragraph("")
        
        p_autor = doc.add_paragraph()
        p_autor.add_run(f"{cliente['nome']}").bold = True
        p_autor.add_run(f", brasileiro(a), {cliente.get('estado_civil', '')}, ")
        p_autor.add_run(f"{cliente.get('profissao', '')}, ")
        p_autor.add_run(f"portador(a) do CPF nº {cpf} e RG nº {cliente.get('rg', '')}, ")
        p_autor.add_run(f"residente e domiciliado(a) em {cliente.get('endereco', '')}, ")
        p_autor.add_run(f"vem, por seu advogado que esta subscreve, propor a presente:")
        
        doc.add_paragraph("")
        
        template = TEMPLATES.get(proc['area'], TEMPLATES['Cível'])
        titulo = doc.add_heading(template['titulo'], level=1)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("")
        doc.add_paragraph("em face de [NOME DO RÉU], pelos seguintes fatos e fundamentos:")
        doc.add_paragraph("")
        
        doc.add_heading("I - DOS FATOS", level=2)
        doc.add_paragraph(proc['resumo'])
        doc.add_paragraph("")
        
        doc.add_heading("II - DO DIREITO", level=2)
        doc.add_paragraph("O presente caso encontra amparo legal nos seguintes dispositivos:")
        doc.add_paragraph("- Constituição Federal de 1988")
        doc.add_paragraph(f"- Legislação específica aplicável à área de {proc['area']}")
        doc.add_paragraph("")
        
        doc.add_heading("III - DOS PEDIDOS", level=2)
        doc.add_paragraph("Diante do exposto, requer-se a Vossa Excelência:")
        
        for i, pedido in enumerate(template['pedidos'], 1):
            doc.add_paragraph(f"{i}. {pedido};")
        
        doc.add_paragraph(f"{len(template['pedidos']) + 1}. A condenação da parte ré ao pagamento de custas processuais e honorários advocatícios;")
        doc.add_paragraph(f"{len(template['pedidos']) + 2}. A produção de todos os meios de prova em direito admitidos.")
        doc.add_paragraph("")
        
        if proc.get('valor_causa'):
            doc.add_paragraph(f"Dá-se à causa o valor de R$ {proc['valor_causa']}.")
        doc.add_paragraph("")
        
        doc.add_paragraph("Termos em que,")
        doc.add_paragraph("Pede deferimento.")
        doc.add_paragraph("")
        doc.add_paragraph(f"{proc.get('comarca', '[Cidade]')}, {datetime.now().strftime('%d de %B de %Y')}.")
        doc.add_paragraph("")
        doc.add_paragraph("_" * 40)
        doc.add_paragraph(f"Advogado(a)")
        doc.add_paragraph("OAB/[UF] nº [NÚMERO]")
        
        nome_limpo = cliente['nome'][:30].replace(' ', '_')
        num_limpo = proc['numero'].replace('/', '-').replace('.', '')
        nome_arquivo = f"PETICAO_{nome_limpo}_{num_limpo}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        doc.save(nome_arquivo)
        abrir_arquivo(nome_arquivo)
        
        msg.showinfo("Sucesso!", f"✅ Petição gerada com sucesso!\n\n📄 Arquivo: {nome_arquivo}")
        
    except Exception as e:
        msg.showerror("Erro", f"Erro ao gerar petição: {str(e)}")

ctk.CTkButton(
    frame_pet,
    text="📄 GERAR PETIÇÃO COMPLETA",
    command=gerar_peticao,
    fg_color=COR_OURO_ESCURO,
    hover_color=COR_OURO,
    height=70,
    width=400,
    font=("Arial", 18, "bold"),
    corner_radius=15
).pack(pady=40)

# ==================== ABA PAINEL ====================
f4 = tab.tab("📊 Painel")
f4.configure(fg_color=COR_FUNDO)

scroll_painel = ctk.CTkScrollableFrame(f4, fg_color="transparent")
scroll_painel.pack(fill="both", expand=True, padx=20, pady=20)

def atualizar_painel():
    for widget in scroll_painel.winfo_children():
        widget.destroy()
    
    ctk.CTkLabel(
        scroll_painel,
        text="PAINEL DE CONTROLE",
        font=("Montserrat", 24, "bold"),
        text_color=COR_OURO
    ).pack(pady=(0, 30))
    
    if not dados["clientes"]:
        ctk.CTkLabel(
            scroll_painel,
            text="📭 Nenhum cliente cadastrado ainda",
            font=("Arial", 16),
            text_color="#888888"
        ).pack(pady=50)
        return
    
    for cpf, cliente in dados["clientes"].items():
        frame_card = ctk.CTkFrame(scroll_painel, fg_color=COR_CARD, corner_radius=15)
        frame_card.pack(fill="x", pady=10, padx=20)
        
        frame_header_card = ctk.CTkFrame(frame_card, fg_color="transparent")
        frame_header_card.pack(fill="x", padx=20, pady=15)
        
        ctk.CTkLabel(
            frame_header_card,
            text=f"👤 {cliente['nome']}",
            font=("Arial", 18, "bold"),
            text_color=COR_OURO
        ).pack(side="left")
        
        def excluir_cliente_func(cpf_del=cpf, nome_del=cliente['nome']):
            if msg.askyesno("Confirmar Exclusão", 
                           f"Deseja realmente excluir o cliente {nome_del} e todos os seus processos?"):
                del dados["clientes"][cpf_del]
                if salvar(dados):
                    msg.showinfo("Sucesso", "Cliente excluído com sucesso!")
                    atualizar_painel()
        
        ctk.CTkButton(
            frame_header_card,
            text="🗑️",
            command=excluir_cliente_func,
            fg_color="#8B0000",
            hover_color="#A52A2A",
            width=40,
            height=30,
            font=("Arial", 16)
        ).pack(side="right")
        
        frame_info_card = ctk.CTkFrame(frame_card, fg_color="transparent")
        frame_info_card.pack(fill="x", padx=20, pady=(0, 10))
        
        info_text = f"📱 {cliente.get('tel', 'N/A')} | 📧 {cliente.get('email', 'N/A')} | 🆔 CPF: {cpf}"
        ctk.CTkLabel(
            frame_info_card,
            text=info_text,
            font=("Arial", 11),
            text_color="#888888"
        ).pack(anchor="w")
        
        if cliente["processos"]:
            frame_processos = ctk.CTkFrame(frame_card, fg_color="#1a1a1a", corner_radius=10)
            frame_processos.pack(fill="x", padx=20, pady=(10, 15))
            
            ctk.CTkLabel(
                frame_processos,
                text=f"⚖️ Processos ({len(cliente['processos'])})",
                font=("Arial", 13, "bold"),
                text_color=COR_OURO
            ).pack(anchor="w", padx=15, pady=(10, 5))
            
            for proc in cliente["processos"]:
                frame_proc_item = ctk.CTkFrame(frame_processos, fg_color="transparent")
                frame_proc_item.pack(fill="x", padx=15, pady=5)
                
                proc_info = f"📋 {proc['numero']} | {proc['area']} | {proc.get('data_cadastro', 'N/A')}"
                
                ctk.CTkLabel(
                    frame_proc_item,
                    text=proc_info,
                    font=("Arial", 11),
                    text_color="#CCCCCC"
                ).pack(side="left")
                
                def excluir_proc_func(cpf_p=cpf, num_p=proc['numero']):
                    if msg.askyesno("Confirmar", f"Excluir o processo {num_p}?"):
                        dados["clientes"][cpf_p]["processos"] = [
                            p for p in dados["clientes"][cpf_p]["processos"] 
                            if p["numero"] != num_p
                        ]
                        if salvar(dados):
                            msg.showinfo("Sucesso", "Processo excluído!")
                            atualizar_painel()
                
                ctk.CTkButton(
                    frame_proc_item,
                    text="❌",
                    command=excluir_proc_func,
                    fg_color="#8B0000",
                    hover_color="#A52A2A",
                    width=30,
                    height=25,
                    font=("Arial", 12)
                ).pack(side="right")
        else:
            ctk.CTkLabel(
                frame_card,
                text="Sem processos cadastrados",
                font=("Arial", 11),
                text_color="#666666"
            ).pack(padx=20, pady=(0, 15))

frame_btn_atualizar = ctk.CTkFrame(scroll_painel, fg_color="transparent")
frame_btn_atualizar.pack(pady=20)

ctk.CTkButton(
    frame_btn_atualizar,
    text="🔄 ATUALIZAR PAINEL",
    command=atualizar_painel,
    fg_color=COR_OURO_ESCURO,
    hover_color=COR_OURO,
    height=50,
    width=250,
    font=("Arial", 14, "bold"),
    corner_radius=10
).pack()

# ==================== ABA BUSCAR ====================
f5 = tab.tab("🔍 Buscar")
f5.configure(fg_color=COR_FUNDO)

frame_busca_container = ctk.CTkFrame(f5, fg_color="transparent")
frame_busca_container.pack(fill="both", expand=True, padx=20, pady=20)

ctk.CTkLabel(
    frame_busca_container,
    text="BUSCA AVANÇADA",
    font=("Montserrat", 24, "bold"),
    text_color=COR_OURO
).pack(pady=(0, 30))

frame_busca = ctk.CTkFrame(frame_busca_container, fg_color=COR_CARD, corner_radius=15)
frame_busca.pack(fill="x", padx=50, pady=(0, 20))

frame_input_busca = ctk.CTkFrame(frame_busca, fg_color="transparent")
frame_input_busca.pack(fill="x", padx=30, pady=20)

ctk.CTkLabel(
    frame_input_busca,
    text="Buscar por Nome, CPF ou Número de Processo:",
    font=("Arial", 14, "bold"),
    text_color=COR_OURO
).pack(anchor="w")

entry_busca = ctk.CTkEntry(
    frame_input_busca,
    width=600,
    height=45,
    font=("Arial", 13),
    placeholder_text="Digite o termo de busca..."
)
entry_busca.pack(side="left", pady=(5, 0))

scroll_resultados = ctk.CTkScrollableFrame(
    frame_busca_container,
    fg_color=COR_CARD,
    corner_radius=15,
    height=500
)
scroll_resultados.pack(fill="both", expand=True, padx=50)

def realizar_busca(*args):
    termo = entry_busca.get().strip().lower()
    
    for widget in scroll_resultados.winfo_children():
        widget.destroy()
    
    if not termo:
        ctk.CTkLabel(
            scroll_resultados,
            text="Digite algo para buscar",
            font=("Arial", 14),
            text_color="#888888"
        ).pack(pady=50)
        return
    
    resultados = []
    
    for cpf, cliente in dados["clientes"].items():
        if (termo in cliente['nome'].lower() or 
            termo in cpf.lower() or
            termo in cliente.get('tel', '').lower() or
            termo in cliente.get('email', '').lower()):
            resultados.append(('cliente', cpf, cliente, None))
        
        for proc in cliente['processos']:
            if (termo in proc['numero'].lower() or 
                termo in proc.get('area', '').lower() or
                termo in proc.get('resumo', '').lower()):
                resultados.append(('processo', cpf, cliente, proc))
    
    if not resultados:
        ctk.CTkLabel(
            scroll_resultados,
            text="😕 Nenhum resultado encontrado",
            font=("Arial", 16),
            text_color="#888888"
        ).pack(pady=50)
        return
    
    ctk.CTkLabel(
        scroll_resultados,
        text=f"✅ {len(resultados)} resultado(s) encontrado(s)",
        font=("Arial", 14, "bold"),
        text_color="#00ff00"
    ).pack(pady=10)
    
    for tipo, cpf, cliente, proc in resultados:
        frame_resultado = ctk.CTkFrame(scroll_resultados, fg_color="#1a1a1a", corner_radius=10)
        frame_resultado.pack(fill="x", padx=20, pady=10)
        
        if tipo == 'cliente':
            ctk.CTkLabel(
                frame_resultado,
                text=f"👤 CLIENTE: {cliente['nome']}",
                font=("Arial", 14, "bold"),
                text_color=COR_OURO
            ).pack(anchor="w", padx=15, pady=(10, 5))
            
            ctk.CTkLabel(
                frame_resultado,
                text=f"CPF: {cpf} | Tel: {cliente.get('tel', 'N/A')} | Email: {cliente.get('email', 'N/A')}",
                font=("Arial", 11),
                text_color="#CCCCCC"
            ).pack(anchor="w", padx=15, pady=(0, 10))
        
        else:
            ctk.CTkLabel(
                frame_resultado,
                text=f"⚖️ PROCESSO: {proc['numero']}",
                font=("Arial", 14, "bold"),
                text_color=COR_OURO
            ).pack(anchor="w", padx=15, pady=(10, 5))
            
            ctk.CTkLabel(
                frame_resultado,
                text=f"Cliente: {cliente['nome']} | Área: {proc['area']} | Data: {proc.get('data_cadastro', 'N/A')}",
                font=("Arial", 11),
                text_color="#CCCCCC"
            ).pack(anchor="w", padx=15, pady=(0, 10))

entry_busca.bind("<KeyRelease>", realizar_busca)

ctk.CTkButton(
    frame_input_busca,
    text="🔍 BUSCAR",
    command=realizar_busca,
    fg_color=COR_OURO_ESCURO,
    hover_color=COR_OURO,
    height=45,
    width=120,
    font=("Arial", 14, "bold"),
    corner_radius=10
).pack(side="left", padx=10, pady=(5, 0))

# ==================== INICIALIZAÇÃO ====================
atualizar_painel()

# Rodapé
frame_footer = ctk.CTkFrame(app, fg_color=COR_CARD, height=40, corner_radius=0)
frame_footer.pack(side="bottom", fill="x")
frame_footer.pack_propagate(False)

ctk.CTkLabel(
    frame_footer,
    text="© 2025 Vanessa Alves Advocacia | Sistema de Gestão Jurídica v2.0 | Todos os direitos reservados",
    font=("Arial", 10),
    text_color="#666666"
).pack(pady=10)

app.mainloop()
